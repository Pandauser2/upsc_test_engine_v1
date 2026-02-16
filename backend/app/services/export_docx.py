"""
Export generated test to .docx: Section 1 Questions, Section 2 Answer key, Section 3 Explanations.
Simple clean format (EXPLORATION §4).
"""
from io import BytesIO
from docx import Document as DocxDocument
from docx.shared import Pt
from sqlalchemy.orm import Session

from app.models.generated_test import GeneratedTest
from app.models.question import Question


def build_docx(test: GeneratedTest, questions: list[Question]) -> BytesIO:
    """Return a BytesIO containing the .docx file (questions, answer key, explanations)."""
    doc = DocxDocument()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    # Section 1: Questions only
    doc.add_heading("Questions", level=0)
    for q in questions:
        p = doc.add_paragraph()
        p.add_run(f"Q{q.sort_order}. ").bold = True
        p.add_run(q.question)
        for opt in ["A", "B", "C", "D"]:
            val = (q.options or {}).get(opt, "")
            doc.add_paragraph(f"  {opt}. {val}", style="List Bullet")

    doc.add_page_break()
    # Section 2: Answer key
    doc.add_heading("Answer key", level=0)
    for q in questions:
        doc.add_paragraph(f"Q{q.sort_order}. {q.correct_option}")

    doc.add_page_break()
    # Section 3: Explanations
    doc.add_heading("Explanations", level=0)
    for q in questions:
        p = doc.add_paragraph()
        p.add_run(f"Q{q.sort_order}. ").bold = True
        p.add_run(q.explanation)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
